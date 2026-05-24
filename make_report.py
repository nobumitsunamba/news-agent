"""
セッションレポートを Word 形式で生成するスクリプト
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ──────────────────────────────────────────────
# ヘルパー関数
# ──────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """表のセル背景色を設定する"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level, color=None):
    """見出しを追加する"""
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_table(doc, headers, rows, header_bg="2E4057", header_fg="FFFFFF"):
    """スタイル付きの表を追加する"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー行
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(*bytes.fromhex(header_fg))
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # データ行
    for r_idx, row in enumerate(rows):
        data_row = table.rows[r_idx + 1]
        bg = "F7F9FC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = data_row.cells[c_idx]
            cell.text = val
            set_cell_bg(cell, bg)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
    return table

def add_code_block(doc, code_text):
    """コードブロック風の段落を追加する"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    p.paragraph_format.left_indent = Cm(0.5)
    # 背景色（段落シェーディング）
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

# ──────────────────────────────────────────────
# ドキュメント生成
# ──────────────────────────────────────────────
doc = Document()

# ページ余白
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# デフォルトフォント
doc.styles['Normal'].font.name = 'メイリオ'
doc.styles['Normal'].font.size = Pt(10.5)

# ══════════════════════════════════════════════
# 表紙
# ══════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('生成AI日報エージェント')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('開発・実装レポート')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x54, 0x78, 0x9B)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = date_p.add_run(f'作成日：{datetime.date.today().strftime("%Y年%m月%d日")}')
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

repo_p = doc.add_paragraph()
repo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = repo_p.add_run('リポジトリ：https://github.com/nobumitsunamba/news-agent')
run4.font.size = Pt(10)
run4.font.color.rgb = RGBColor(0x08, 0x76, 0xC8)

doc.add_page_break()

# ══════════════════════════════════════════════
# 1. プロジェクト概要
# ══════════════════════════════════════════════
add_heading(doc, '1. プロジェクト概要', 1, '2E4057')

doc.add_paragraph(
    '本プロジェクトは、生成AI関連ニュースを毎朝自動収集・要約してGmailで送信する'
    'Pythonスクリプトを開発・GitHub Actionsに自動実行環境を構築したものです。'
)

add_heading(doc, '1.1 目的', 2)
items = [
    '複数のRSSフィードから前日分の生成AI関連ニュースを自動収集する',
    'キーワードフィルタリングにより無関係な記事を除外する',
    'Claude APIを活用して重複排除・重要度ソート・日本語3行要約を行う',
    'Gmailで毎朝7時（JST）に自動送信する',
    'GitHub Actionsを利用してサーバー不要で完全自動化する',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].font.size = Pt(10.5)

add_heading(doc, '1.2 開発環境', 2)
add_table(doc,
    ['項目', '内容'],
    [
        ['言語', 'Python 3.12'],
        ['AIモデル', 'claude-sonnet-4-6（Anthropic）'],
        ['実行環境', 'GitHub Actions（ubuntu-latest）'],
        ['メール送信', 'Gmail SMTP（ポート465 / SSL）'],
        ['スケジューラ', 'GitHub Actions cron（UTC 22:00 = JST 07:00）'],
    ]
)

doc.add_paragraph()

# ══════════════════════════════════════════════
# 2. システム構成
# ══════════════════════════════════════════════
add_heading(doc, '2. システム構成', 1, '2E4057')

add_heading(doc, '2.1 ファイル構成', 2)
add_code_block(doc,
    'news-agent/\n'
    '├── news_agent.py                   # メインスクリプト\n'
    '├── requirements.txt                # Pythonパッケージ定義\n'
    '├── .gitignore                      # .env を管理対象外に設定\n'
    '├── README.md                       # セットアップ・運用手順書\n'
    '└── .github/\n'
    '    └── workflows/\n'
    '        └── news_agent.yml          # GitHub Actions ワークフロー'
)
doc.add_paragraph()

add_heading(doc, '2.2 処理フロー', 2)
add_table(doc,
    ['ステップ', '処理内容', '実装箇所'],
    [
        ['① RSS収集', '4つのRSSフィードから前日分の記事を取得', 'collect_all_articles()'],
        ['② フィルタリング', '24種類のキーワードで生成AI関連記事に絞り込み', 'filter_articles()'],
        ['③ Claude API処理', '重複排除・重要度ソート・3行日本語要約', 'summarize_with_claude()'],
        ['④ メール送信', 'Gmail SMTPでプレーンテキスト＋HTML送信', 'send_email()'],
    ]
)
doc.add_paragraph()

add_heading(doc, '2.3 収集対象RSSフィード', 2)
add_table(doc,
    ['フィード', '内容'],
    [
        ['Google News（生成AI）', '生成AI関連の最新ニュース（日本語）'],
        ['Google News（ChatGPT/Claude/Gemini/LLM）', '主要AIサービス関連ニュース（日本語）'],
        ['ITmedia NEWS', 'IT専門メディア'],
        ['TechCrunch Japan', 'スタートアップ・テック系メディア'],
    ]
)
doc.add_paragraph()

add_heading(doc, '2.4 キーワードフィルタリング', 2)
doc.add_paragraph(
    '以下の24種類のキーワードのいずれかをタイトルまたはサマリに含む記事のみを対象とします。'
)
kws = ('生成AI、ChatGPT、Claude、Gemini、LLM、大規模言語モデル、Anthropic、OpenAI、'
       'Google DeepMind、GPT、AI、人工知能、機械学習、ディープラーニング、深層学習、'
       'Grok、Llama、Mistral、Copilot、Sora、Stable Diffusion、画像生成、音声生成、自然言語処理')
p = doc.add_paragraph(kws)
p.paragraph_format.left_indent = Cm(0.5)
p.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

# ══════════════════════════════════════════════
# 3. 実装詳細
# ══════════════════════════════════════════════
add_heading(doc, '3. 実装詳細', 1, '2E4057')

add_heading(doc, '3.1 Claude APIへのプロンプト設計', 2)
doc.add_paragraph('Claude APIには1回のAPIコールで以下の3つの処理を依頼します。')
add_table(doc,
    ['処理', '内容'],
    [
        ['重複排除', '同一・類似記事を統合し代表1件のみ残す'],
        ['重要度ソート', '影響度・話題性の高い順に並び替える'],
        ['3行要約', '各記事を3行以内の日本語で要約する'],
    ]
)
doc.add_paragraph()

add_heading(doc, '3.2 メール仕様', 2)
add_table(doc,
    ['項目', '内容'],
    [
        ['件名', '【生成AI日報】YYYY年MM月DD日'],
        ['送信方式', 'Gmail SMTP SSL（ポート465）'],
        ['認証', 'Googleアプリパスワード（16桁）'],
        ['本文形式', 'プレーンテキスト＋HTML（マルチパート）'],
        ['送信タイミング', '毎朝7:00 JST（GitHub Actions cron）'],
    ]
)
doc.add_paragraph()

add_heading(doc, '3.3 GitHub Actions ワークフロー', 2)
add_code_block(doc,
    'on:\n'
    '  schedule:\n'
    "    - cron: '0 22 * * *'   # UTC 22:00 = JST 07:00\n"
    '  workflow_dispatch:        # 手動実行も可能\n\n'
    'jobs:\n'
    '  send-news-digest:\n'
    '    runs-on: ubuntu-latest\n'
    '    timeout-minutes: 15'
)
doc.add_paragraph()

add_heading(doc, '3.4 必要な環境変数（GitHub Secrets）', 2)
add_table(doc,
    ['Secret名', '説明', '取得先'],
    [
        ['ANTHROPIC_API_KEY', 'Claude API キー', 'console.anthropic.com'],
        ['GMAIL_ADDRESS', '送信元Gmailアドレス', 'Googleアカウント'],
        ['GMAIL_APP_PASSWORD', 'Gmailアプリパスワード（16桁）', 'Googleアカウント → セキュリティ'],
        ['TO_EMAIL', '受信先メールアドレス', '任意'],
    ]
)
doc.add_paragraph()

# ══════════════════════════════════════════════
# 4. 開発経緯・トラブルシューティング
# ══════════════════════════════════════════════
add_heading(doc, '4. 開発経緯・発生したトラブルと対処', 1, '2E4057')

add_table(doc,
    ['発生事象', '原因', '対処'],
    [
        [
            'GitHub Actions実行時に\n404エラーが発生',
            'モデル名 claude-sonnet-4-20250514 が\n廃止予定かつAPIで404を返した',
            'claude-sonnet-4-6 に変更\n（PR #2 でmainにマージ）',
        ],
    ]
)
doc.add_paragraph()

add_heading(doc, '4.1 最終動作確認結果', 2)
add_table(doc,
    ['確認項目', '結果', '備考'],
    [
        ['RSSフィード収集', '✅ 成功', '62件取得（4フィード合計）'],
        ['キーワードフィルタリング', '✅ 成功', '48件に絞り込み'],
        ['Claude API要約', '✅ 成功', 'claude-sonnet-4-6 使用'],
        ['Gmail送信', '✅ 成功', 'メール受信を確認'],
    ]
)
doc.add_paragraph()

# ══════════════════════════════════════════════
# 5. 運用・カスタマイズ方法
# ══════════════════════════════════════════════
add_heading(doc, '5. 運用・カスタマイズ方法', 1, '2E4057')

add_table(doc,
    ['やりたいこと', '変更箇所', '変更方法'],
    [
        ['RSSフィードを追加', 'news_agent.py', 'RSS_FEEDS リストにURLを追加'],
        ['キーワードを追加・削除', 'news_agent.py', 'AI_KEYWORDS リストを編集'],
        ['送信時刻を変更', 'news_agent.yml', 'cron の値を変更（UTC基準）'],
        ['モデルを変更（コスト削減）', 'news_agent.py', 'CLAUDE_MODEL を haiku に変更'],
        ['手動実行', 'GitHub Actions', 'Actionsタブ → Run workflow'],
    ]
)
doc.add_paragraph()

add_heading(doc, '5.1 コスト目安', 2)
add_table(doc,
    ['モデル', '1回あたりのコスト目安', '月額目安（毎日実行）'],
    [
        ['claude-sonnet-4-6', '$0.05〜$0.20（7〜30円）', '約200〜900円'],
        ['claude-haiku-4-5', '$0.01〜$0.05（2〜7円）', '約60〜210円'],
    ]
)
doc.add_paragraph()

# ══════════════════════════════════════════════
# 6. 成果物一覧
# ══════════════════════════════════════════════
add_heading(doc, '6. 成果物一覧', 1, '2E4057')

add_table(doc,
    ['ファイル', '説明'],
    [
        ['news_agent.py', 'メインスクリプト（RSS収集・フィルタ・Claude要約・メール送信）'],
        ['requirements.txt', 'Pythonパッケージ定義（anthropic / feedparser / python-dotenv）'],
        ['.github/workflows/news_agent.yml', 'GitHub Actions ワークフロー（毎朝JST 7:00自動実行）'],
        ['.gitignore', '.envファイルをGit管理外に設定'],
        ['README.md', 'GitHub Secrets設定手順・セットアップガイド'],
    ]
)
doc.add_paragraph()

add_heading(doc, '6.1 GitHubプルリクエスト履歴', 2)
add_table(doc,
    ['PR番号', 'タイトル', 'ステータス'],
    [
        ['#1', 'Add automated AI news digest agent with Claude summarization', 'Merged'],
        ['#2', 'fix: Claude モデルを claude-sonnet-4-6 に更新', 'Open'],
    ]
)

doc.add_paragraph()
doc.add_paragraph()

# フッター
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = footer_p.add_run('── 以上 ──')
run_f.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run_f.font.size = Pt(9)

# 保存
output_path = '/home/user/news-agent/生成AI日報エージェント_開発レポート.docx'
doc.save(output_path)
print(f'✅ 保存完了: {output_path}')
